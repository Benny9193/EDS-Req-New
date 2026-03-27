#!/usr/bin/env python3
"""
Generate formatted DOCX files from Markdown documentation.

Creates visually polished Word documents with proper styling for:
- Headers (H1, H2, H3)
- Tables with borders and shading
- Code blocks with monospace font and background
- Bulleted and numbered lists
- Bold and italic text
- Horizontal rules

Usage:
    python scripts/generate_docx.py
    python scripts/generate_docx.py --file docs/SCHEMA.md
    python scripts/generate_docx.py --output-dir docs/docx
"""

import re
import os
import sys
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# EDS Brand Colors
EDS_NAVY = RGBColor(0x1C, 0x1A, 0x83)
EDS_DARK_BLUE = RGBColor(0x1A, 0x36, 0x5D)
EDS_RED = RGBColor(0xB7, 0x0C, 0x0D)
CODE_BG = RGBColor(0xF5, 0xF5, 0xF5)


def create_styled_document():
    """Create a new document with custom styles."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title style
    if 'Doc Title' not in [s.name for s in doc.styles]:
        title_style = doc.styles.add_style('Doc Title', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Calibri Light'
        title_style.font.size = Pt(28)
        title_style.font.color.rgb = EDS_NAVY
        title_style.font.bold = True
        title_style.paragraph_format.space_after = Pt(12)
        title_style.paragraph_format.space_before = Pt(0)

    # Heading 1
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri Light'
    h1.font.size = Pt(20)
    h1.font.color.rgb = EDS_NAVY
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)

    # Heading 2
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Calibri Light'
    h2.font.size = Pt(16)
    h2.font.color.rgb = EDS_DARK_BLUE
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(4)

    # Heading 3
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(13)
    h3.font.color.rgb = EDS_DARK_BLUE
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(2)

    # Code style
    if 'Code' not in [s.name for s in doc.styles]:
        code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Consolas'
        code_style.font.size = Pt(9)
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)
        code_style.paragraph_format.left_indent = Inches(0.25)

    # Code Char style for inline code
    if 'Code Char' not in [s.name for s in doc.styles]:
        code_char = doc.styles.add_style('Code Char', WD_STYLE_TYPE.CHARACTER)
        code_char.font.name = 'Consolas'
        code_char.font.size = Pt(10)

    return doc


def add_horizontal_rule(doc):
    """Add a horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    # Add bottom border to paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        r'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
        r'</w:pBdr>'
    )
    pPr.append(pBdr)


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, rows, has_header=True):
    """Add a formatted table."""
    if not rows:
        return

    # Create table
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set column widths based on content
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = cell_text.strip()

            # Style header row
            if i == 0 and has_header:
                set_cell_shading(cell, "1C1A83")  # EDS Navy
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.size = Pt(10)
            else:
                # Alternate row shading
                if i % 2 == 0:
                    set_cell_shading(cell, "F5F5F5")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

    # Add space after table
    doc.add_paragraph()


def add_code_block(doc, code, language=''):
    """Add a formatted code block."""
    # Add a paragraph with code styling
    for line in code.split('\n'):
        p = doc.add_paragraph(style='Code')
        run = p.add_run(line if line else ' ')  # Keep empty lines
        run.font.name = 'Consolas'
        run.font.size = Pt(9)


def parse_inline_formatting(paragraph, text):
    """Parse and apply inline formatting (bold, italic, code)."""
    # Pattern for inline code, bold, italic
    pattern = r'(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*|__[^_]+__|_[^_]+_|[^`*_]+)'

    parts = re.findall(pattern, text)

    for part in parts:
        if not part:
            continue

        if part.startswith('`') and part.endswith('`'):
            # Inline code
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        elif (part.startswith('**') and part.endswith('**')) or \
             (part.startswith('__') and part.endswith('__')):
            # Bold
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif (part.startswith('*') and part.endswith('*')) or \
             (part.startswith('_') and part.endswith('_')):
            # Italic
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            # Regular text
            paragraph.add_run(part)


def convert_markdown_to_docx(md_content, doc):
    """Convert markdown content to DOCX."""
    lines = md_content.split('\n')
    i = 0

    in_code_block = False
    code_content = []
    code_language = ''

    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]

        # Code block handling
        if line.startswith('```'):
            if in_code_block:
                # End code block
                add_code_block(doc, '\n'.join(code_content), code_language)
                code_content = []
                code_language = ''
                in_code_block = False
            else:
                # Start code block
                in_code_block = True
                code_language = line[3:].strip()
            i += 1
            continue

        if in_code_block:
            code_content.append(line)
            i += 1
            continue

        # Table handling
        if '|' in line and not line.startswith('```'):
            # Check if it's a table row
            cells = [c.strip() for c in line.split('|')]
            cells = [c for c in cells if c or cells.index(c) not in [0, len(cells)-1]]

            if cells and not all(c.replace('-', '').replace(':', '') == '' for c in cells):
                # Not a separator row
                if not in_table:
                    in_table = True
                    table_rows = []
                table_rows.append(cells)
            elif in_table and all(c.replace('-', '').replace(':', '').replace(' ', '') == '' for c in cells):
                # Separator row, skip
                pass
            i += 1
            continue
        elif in_table:
            # End of table
            add_table(doc, table_rows)
            table_rows = []
            in_table = False

        # Horizontal rule
        if line.strip() in ['---', '***', '___']:
            add_horizontal_rule(doc)
            i += 1
            continue

        # Headers
        if line.startswith('# '):
            p = doc.add_paragraph(line[2:].strip(), style='Doc Title')
            i += 1
            continue
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=3)
            i += 1
            continue

        # Bulleted list
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            parse_inline_formatting(p, text)
            i += 1
            continue

        # Numbered list
        if re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            p = doc.add_paragraph(style='List Number')
            parse_inline_formatting(p, text)
            i += 1
            continue

        # Regular paragraph
        if line.strip():
            p = doc.add_paragraph()
            parse_inline_formatting(p, line)

        i += 1

    # Handle any remaining table
    if in_table and table_rows:
        add_table(doc, table_rows)


def add_cover_page(doc, title, subtitle=''):
    """Add a cover page."""
    # Add space at top
    for _ in range(8):
        doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.name = 'Calibri Light'
    run.font.size = Pt(36)
    run.font.color.rgb = EDS_NAVY
    run.bold = True

    # Subtitle
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        run.font.name = 'Calibri'
        run.font.size = Pt(16)
        run.font.color.rgb = EDS_DARK_BLUE

    # Add space
    for _ in range(4):
        doc.add_paragraph()

    # Add line
    add_horizontal_rule(doc)

    # Footer info
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('EDS Universal Requisition System')
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.color.rgb = EDS_DARK_BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Page break
    doc.add_page_break()


def get_title_from_content(content):
    """Extract title from markdown content."""
    lines = content.split('\n')
    for line in lines:
        if line.startswith('# '):
            return line[2:].strip()
    return 'Documentation'


def get_subtitle_from_filename(filename):
    """Generate subtitle from filename."""
    subtitles = {
        # Main Documentation
        'DEVELOPMENT': 'Developer Setup Guide',
        'DEPLOYMENT': 'Deployment & Docker Guide',
        'TESTING': 'Testing Guide',
        'CONFIGURATION': 'Configuration Reference',
        'TROUBLESHOOTING': 'Troubleshooting Guide',
        'ARCHITECTURE': 'System Architecture',
        'SCHEMA': 'Database Schema Reference',
        'API_REFERENCE': 'API Documentation',
        'UNIVERSAL_REQUISITION': 'Frontend User Guide',
        'AGENT_CLI': 'AI Agent & CLI Guide',
        'README': 'Project Overview',
        # Database Schema & Structure
        'EDS_SUMMARY': 'Database Summary',
        'EDS_DATA_DICTIONARY': 'Complete Data Dictionary',
        'EDS_ERD': 'Entity Relationship Diagrams',
        # Business Logic & Domains
        'EDS_BUSINESS_DOMAINS': 'Business Domain Analysis',
        'EDS_BUSINESS_WORKFLOWS': 'Business Workflow Documentation',
        'EDS_STATUS_CODES': 'Status Codes Reference',
        'EDS_DATA_OWNERSHIP': 'Data Ownership Guide',
        'EDS_ACCESS_CONTROL': 'Access Control Documentation',
        # Stored Procedures
        'EDS_STORED_PROCEDURES': 'Stored Procedures Reference',
        'EDS_PROCEDURES_GUIDE': 'Procedures Usage Guide',
        'EDS_ROOT_PROCEDURES': 'Root Procedures Documentation',
        'EDS_PROCEDURE_DEPENDENCIES': 'Procedure Dependencies',
        'EDS_SP_DEPENDENCIES': 'SP Dependencies Analysis',
        'EDS_RECURSIVE_PROCEDURES': 'Recursive Procedures',
        'EDS_CIRCULAR_DEPS': 'Circular Dependencies Analysis',
        'EDS_INFINITE_LOOP_ANALYSIS': 'Infinite Loop Analysis',
        # Views
        'EDS_VIEWS': 'Views Reference',
        'EDS_VIEWS_GUIDE': 'Views Usage Guide',
        # Triggers & Indexes
        'EDS_TRIGGERS': 'Triggers Documentation',
        'EDS_INDEXES': 'Index Documentation',
        # Archive & ETL
        'EDS_ARCHIVE_ANALYSIS': 'Archive Analysis',
        'EDS_ARCHIVE_STRATEGY': 'Archive Strategy Guide',
        'EDS_ETL_INTEGRATIONS': 'ETL Integration Guide',
    }

    name = Path(filename).stem.upper()
    return subtitles.get(name, '')


def convert_file(input_path, output_path):
    """Convert a single markdown file to DOCX."""
    print(f"  Converting: {input_path}")

    # Read markdown content
    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create document
    doc = create_styled_document()

    # Add cover page
    title = get_title_from_content(content)
    subtitle = get_subtitle_from_filename(input_path)
    add_cover_page(doc, title, subtitle)

    # Convert content
    convert_markdown_to_docx(content, doc)

    # Save
    doc.save(output_path)
    print(f"  Created: {output_path}")


def main():
    """Main function."""
    import argparse

    parser = argparse.ArgumentParser(description='Convert Markdown to DOCX')
    parser.add_argument('--file', help='Single file to convert')
    parser.add_argument('--output-dir', default='docs/docx', help='Output directory')
    args = parser.parse_args()

    # Find project root
    script_dir = Path(__file__).parent
    project_root = script_dir.parent
    docs_dir = project_root / 'docs'
    output_dir = project_root / args.output_dir

    # Create output directory
    output_dir.mkdir(parents=True, exist_ok=True)

    print(f"\n{'='*60}")
    print("EDS Documentation - DOCX Generator")
    print(f"{'='*60}")
    print(f"Output directory: {output_dir}\n")

    if args.file:
        # Convert single file
        input_path = Path(args.file)
        if not input_path.exists():
            input_path = project_root / args.file

        output_path = output_dir / f"{input_path.stem}.docx"
        convert_file(input_path, output_path)
    else:
        # Convert all documentation files (34 total)
        files_to_convert = [
            # Main Documentation
            'README.md',
            'ARCHITECTURE.md',
            'DEVELOPMENT.md',
            'DEPLOYMENT.md',
            'CONFIGURATION.md',
            'TESTING.md',
            'TROUBLESHOOTING.md',
            'API_REFERENCE.md',
            'UNIVERSAL_REQUISITION.md',
            'AGENT_CLI.md',
            # Database Schema & Structure
            'SCHEMA.md',
            'EDS_SUMMARY.md',
            'EDS_DATA_DICTIONARY.md',
            'EDS_ERD.md',
            # Business Logic & Domains
            'EDS_BUSINESS_DOMAINS.md',
            'EDS_BUSINESS_WORKFLOWS.md',
            'EDS_STATUS_CODES.md',
            'EDS_DATA_OWNERSHIP.md',
            'EDS_ACCESS_CONTROL.md',
            # Stored Procedures
            'EDS_STORED_PROCEDURES.md',
            'EDS_PROCEDURES_GUIDE.md',
            'EDS_ROOT_PROCEDURES.md',
            'EDS_PROCEDURE_DEPENDENCIES.md',
            'EDS_SP_DEPENDENCIES.md',
            'EDS_RECURSIVE_PROCEDURES.md',
            'EDS_CIRCULAR_DEPS.md',
            'EDS_INFINITE_LOOP_ANALYSIS.md',
            # Views
            'EDS_VIEWS.md',
            'EDS_VIEWS_GUIDE.md',
            # Triggers & Indexes
            'EDS_TRIGGERS.md',
            'EDS_INDEXES.md',
            # Archive & ETL
            'EDS_ARCHIVE_ANALYSIS.md',
            'EDS_ARCHIVE_STRATEGY.md',
            'EDS_ETL_INTEGRATIONS.md',
        ]

        converted = 0
        for filename in files_to_convert:
            input_path = docs_dir / filename
            if input_path.exists():
                output_path = output_dir / f"{input_path.stem}.docx"
                try:
                    convert_file(input_path, output_path)
                    converted += 1
                except Exception as e:
                    print(f"  ERROR converting {filename}: {e}")
            else:
                print(f"  Skipping (not found): {filename}")

        print(f"\n{'='*60}")
        print(f"Conversion complete: {converted} files created")
        print(f"Output location: {output_dir}")
        print(f"{'='*60}\n")


if __name__ == '__main__':
    main()
