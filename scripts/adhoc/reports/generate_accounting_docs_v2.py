"""Generate enhanced DOCX documentation for each EDS accounting system.
Incorporates deep data: stored proc source, yearly trends, approval workflows,
account code patterns, PO layouts, transaction types, status flows."""
import json
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree
from collections import defaultdict
from datetime import datetime
import pyodbc
import decimal

# Database connection
conn_str = (
    'DRIVER={ODBC Driver 17 for SQL Server};'
    'SERVER=eds-sqlserver.eastus2.cloudapp.azure.com;'
    'DATABASE=EDS;'
    'UID=EDSAdmin;'
    'PWD=Consultant~!;'
    'TrustServerCertificate=yes;'
)

OUTPUT_DIR = 'docs/accounting'
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Load deep data
print("Loading deep data...")
with open('output/accounting_deep_data.json', 'r') as f:
    deep = json.load(f)

conn = pyodbc.connect(conn_str)
cursor = conn.cursor()

# ============================================================
# UTILITY FUNCTIONS
# ============================================================

def set_cell_shading(cell, color_hex):
    shading = etree.SubElement(cell._tc.get_or_add_tcPr(), qn('w:shd'))
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')


def add_styled_table(doc, headers, rows, col_widths=None):
    if not rows:
        doc.add_paragraph('No data available.')
        return None
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '1c1a83')

    for r_idx, row_data in enumerate(rows):
        for c_idx, value in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(value) if value is not None else ''
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'f0f0f8')
    return table


def add_title_page(doc, title, subtitle=''):
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(28, 26, 131)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(subtitle)
        run2.font.size = Pt(14)
        run2.font.color.rgb = RGBColor(74, 72, 144)
    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run('Educational Data Services (EDS)')
    run3.font.size = Pt(12)
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
    run4.font.size = Pt(10)
    run4.font.color.rgb = RGBColor(128, 128, 128)
    doc.add_page_break()


def add_code_block(doc, code_text, max_lines=60):
    """Add a formatted code block to the document."""
    lines = code_text.strip().split('\n')
    if len(lines) > max_lines:
        lines = lines[:max_lines] + [f'... ({len(code_text.split(chr(10))) - max_lines} more lines)']
    text = '\n'.join(lines)
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(51, 51, 51)
    # Light gray background for the paragraph
    pPr = p._p.get_or_add_pPr()
    shd = etree.SubElement(pPr, qn('w:shd'))
    shd.set(qn('w:fill'), 'F5F5F5')
    shd.set(qn('w:val'), 'clear')


def add_info_box(doc, text, color='4a4890'):
    """Add an info/note box."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = RGBColor(int(color[:2], 16), int(color[2:4], 16), int(color[4:6], 16))


# ============================================================
# GATHER LIVE DATA
# ============================================================
print("Gathering live data from database...")

cursor.execute("SELECT * FROM dbo.AccountingFormats ORDER BY AccountingFormatId")
fmt_cols = [d[0] for d in cursor.description]
all_formats = [dict(zip(fmt_cols, row)) for row in cursor.fetchall()]
fmt_by_id = {f['AccountingFormatId']: f for f in all_formats}

cursor.execute("""
    SELECT d.DistrictId, d.DistrictCode, d.Name, d.City, d.State, d.County,
           d.AccountingFormatId, d.RequireAccounts, d.CurrentBudgetOnly,
           d.AccountSeparator, d.AccountingDistrictCode, d.AccountingSystemOptions,
           d.RequiredApprovalLevel,
           af.Description AS AccountingSystem, af.ShortName,
           af.LocationCodeRequired, af.VendorBidNumberRequired,
           af.VendorBidCommentsRequired, af.UsersDistrictAccountingCodeRequired,
           af.IncidentalOrdersSupported, af.DetailedFormat,
           (SELECT COUNT(*) FROM dbo.School s WHERE s.DistrictId = d.DistrictId) AS SchoolCount,
           (SELECT COUNT(*) FROM dbo.Accounts a WHERE a.DistrictId = d.DistrictId AND a.Active = 1) AS AccountCount
    FROM dbo.District d
    LEFT JOIN dbo.AccountingFormats af ON d.AccountingFormatId = af.AccountingFormatId
    WHERE d.Active = 1
    ORDER BY af.Description, d.Name
""")
dist_cols = [d[0] for d in cursor.description]
all_districts = [dict(zip(dist_cols, row)) for row in cursor.fetchall()]

SYSTEM_FAMILIES = {
    'Finance Manager': lambda d: 'Finance Manager' in (d.get('AccountingSystem') or ''),
    'Systems 3000': lambda d: 'System' in (d.get('AccountingSystem') or '') and '3000' in (d.get('AccountingSystem') or ''),
    'CSI': lambda d: 'CSI' in (d.get('AccountingSystem') or ''),
    'Genesis': lambda d: 'Genesis' in (d.get('AccountingSystem') or ''),
    'WinCap': lambda d: 'Wincap' in (d.get('AccountingSystem') or '') or 'WinCap' in (d.get('ShortName') or ''),
    'CDK': lambda d: 'CDK' in (d.get('AccountingSystem') or ''),
    'Edu-Met': lambda d: 'Edu-Met' in (d.get('AccountingSystem') or '') or 'EMapp' in (d.get('AccountingSystem') or ''),
    'Sage/Alio': lambda d: 'Sage' in (d.get('AccountingSystem') or '') or 'Alio' in (d.get('AccountingSystem') or ''),
    'Edmunds': lambda d: 'Edmunds' in (d.get('AccountingSystem') or ''),
    'Infinite Visions': lambda d: 'Infinite' in (d.get('AccountingSystem') or ''),
    'No Interface': lambda d: 'No Accounting' in (d.get('AccountingSystem') or ''),
    'Pentamation': lambda d: 'Pentamation' in (d.get('AccountingSystem') or ''),
    'APECS': lambda d: 'APECS' in (d.get('AccountingSystem') or ''),
    'Keystone': lambda d: 'Keystone' in (d.get('AccountingSystem') or ''),
    'Asbury Park': lambda d: 'Asbury' in (d.get('AccountingSystem') or ''),
}

districts_by_family = defaultdict(list)
for d in all_districts:
    matched = False
    for family, predicate in SYSTEM_FAMILIES.items():
        if predicate(d):
            districts_by_family[family].append(d)
            matched = True
            break
    if not matched:
        sys_name = d.get('AccountingSystem') or 'Unassigned'
        districts_by_family[sys_name].append(d)

# Account codes
cursor.execute("""
    SELECT d.DistrictId, af.Description AS SystemName, d.Name AS DistrictName,
           a.Code, a.Description AS AcctDesc
    FROM dbo.Accounts a
    JOIN dbo.District d ON a.DistrictId = d.DistrictId
    LEFT JOIN dbo.AccountingFormats af ON d.AccountingFormatId = af.AccountingFormatId
    WHERE a.Active = 1 AND d.Active = 1
""")
acct_cols = [d[0] for d in cursor.description]
all_accounts = [dict(zip(acct_cols, row)) for row in cursor.fetchall()]
accounts_by_district = defaultdict(list)
for a in all_accounts:
    accounts_by_district[a['DistrictId']].append(a)

# Budget data
cursor.execute("""
    SELECT d.DistrictId, d.Name AS DistrictName, af.Description AS SystemName,
           b.Name AS BudgetName, b.Active, b.StartDate, b.EndDate,
           COUNT(ba.BudgetAccountId) AS AccountCount,
           SUM(ba.BudgetAmount) AS TotalBudget
    FROM dbo.Budgets b
    JOIN dbo.District d ON b.DistrictId = d.DistrictId
    LEFT JOIN dbo.AccountingFormats af ON d.AccountingFormatId = af.AccountingFormatId
    LEFT JOIN dbo.BudgetAccounts ba ON b.BudgetId = ba.BudgetId
    WHERE d.Active = 1
    GROUP BY d.DistrictId, d.Name, af.Description, b.Name, b.Active, b.StartDate, b.EndDate
    ORDER BY d.Name, b.StartDate DESC
""")
bdg_cols = [d[0] for d in cursor.description]
all_budgets = [dict(zip(bdg_cols, row)) for row in cursor.fetchall()]
budgets_by_district = defaultdict(list)
for b in all_budgets:
    budgets_by_district[b['DistrictId']].append(b)

# Requisition volumes by format description
cursor.execute("""
    SELECT af.Description AS SystemName,
           COUNT(r.RequisitionId) AS ReqCount,
           MIN(r.DateEntered) AS EarliestReq,
           MAX(r.DateEntered) AS LatestReq
    FROM dbo.Requisitions r
    JOIN dbo.School s ON r.SchoolId = s.SchoolId
    JOIN dbo.District d ON s.DistrictId = d.DistrictId
    LEFT JOIN dbo.AccountingFormats af ON d.AccountingFormatId = af.AccountingFormatId
    WHERE d.Active = 1
    GROUP BY af.Description
""")
req_cols = [d[0] for d in cursor.description]
req_volumes = {row[0]: dict(zip(req_cols, row)) for row in cursor.fetchall()}

cursor.execute("""
    SELECT af.Description AS SystemName, COUNT(p.POId) AS POCount
    FROM dbo.PO p
    JOIN dbo.Requisitions r ON p.RequisitionId = r.RequisitionId
    JOIN dbo.School s ON r.SchoolId = s.SchoolId
    JOIN dbo.District d ON s.DistrictId = d.DistrictId
    LEFT JOIN dbo.AccountingFormats af ON d.AccountingFormatId = af.AccountingFormatId
    WHERE d.Active = 1
    GROUP BY af.Description
""")
po_cols = [d[0] for d in cursor.description]
po_volumes = {row[0]: dict(zip(po_cols, row)) for row in cursor.fetchall()}

# User fields
cursor.execute("""
    SELECT auf.AccountingFormatId, af.Description AS FormatName,
           auf.DistrictId, d.Name AS DistrictName,
           auf.FieldPos, auf.FieldName, auf.RequiredField
    FROM dbo.AccountingUserFields auf
    JOIN dbo.AccountingFormats af ON auf.AccountingFormatId = af.AccountingFormatId
    JOIN dbo.District d ON auf.DistrictId = d.DistrictId
    ORDER BY auf.AccountingFormatId, auf.DistrictId, auf.FieldPos
""")
uf_cols = [d[0] for d in cursor.description]
all_user_fields = [dict(zip(uf_cols, row)) for row in cursor.fetchall()]
user_fields_by_format = defaultdict(list)
for uf in all_user_fields:
    user_fields_by_format[uf['AccountingFormatId']].append(uf)

print(f"Data gathered: {len(all_districts)} districts, {len(all_accounts)} accounts")

# ============================================================
# DEEP DATA HELPERS
# ============================================================

# Index yearly trends by system name
yearly_reqs_by_system = defaultdict(list)
for entry in deep.get('yearly_reqs', []):
    yearly_reqs_by_system[entry.get('SystemName', '')].append(entry)

yearly_pos_by_system = defaultdict(list)
for entry in deep.get('yearly_pos', []):
    yearly_pos_by_system[entry.get('SystemName', '')].append(entry)

# Status distribution by system
status_by_system = defaultdict(list)
for entry in deep.get('status_by_system', []):
    status_by_system[entry.get('SystemName', '')].append(entry)

# Code patterns by system
code_patterns_by_system = defaultdict(list)
for entry in deep.get('code_patterns', []):
    code_patterns_by_system[entry.get('SystemName', '')].append(entry)

# Approval by system
approval_by_system = defaultdict(list)
for entry in deep.get('approval_by_system', []):
    approval_by_system[entry.get('SystemName', '')].append(entry)

# PO layout assignments
po_layout_by_system = defaultdict(list)
for entry in deep.get('po_layout_by_system', []):
    po_layout_by_system[entry.get('SystemName', '')].append(entry)

# All statuses
all_statuses = deep.get('statuses', [])
status_map = {s.get('StatusId'): s.get('Name', '') for s in all_statuses}

# Proc sources
proc_sources = deep.get('proc_sources', {})
func_sources = deep.get('func_sources', {})

# Key stored procedures grouped by function
PROC_GROUPS = {
    'Account Management': [
        'sp_FA_AddUpdateAccountCode', 'sp_FA_AvailableAccounts', 'sp_FA_DeleteAccount',
        'sp_FA_SetBudgetAccount', 'sp_FA_SetUserAccount',
        'sp_RefreshAccounts', 'sp_MergeAccounts',
    ],
    'Requisition Lifecycle': [
        'sp_CreateNewRequisition', 'sp_SubmitRequisition', 'sp_SubmitRequisitionNew',
        'sp_DeleteRequisition', 'sp_HoldRequisition', 'sp_CanDeleteRequisition',
        'sp_FA_UpdateRequisitionStatus', 'sp_FA_DeleteRequisition',
        'sp_GetUserRequisitions',
    ],
    'Budget & Fiscal Year': [
        'sp_SetBudgetYear', 'sp_SetDistrictAndBudgetYear', 'sp_ResetDistrictAccountingYear',
        'sp_CopyBudgetAmounts', 'sp_MasterBudgetBook',
        'usp_BringAccountsForward', 'usp_UpdateBudgets',
    ],
    'PO Processing': [
        'sp_FA_RequisitionsForPurchaseOrderModal', 'sp_FA_RequisitionsTotals',
        'usp_POPrintExport',
    ],
    'Notes & Communication': [
        'sp_FA_SaveRequisitionNote', 'sp_FA_SaveRequisitionNoteEmails',
    ],
    'User Administration': [
        'sp_CCAccountMaint', 'sp_CCUpdateUserAccounts', 'sp_CCUserAccountMaint',
        'sp_FA_AttemptLogin',
    ],
    'Public Access': [
        'sp_PAAccounts', 'sp_PABudgets', 'sp_PARequisitions',
    ],
}

FUNC_GROUPS = {
    'Account Functions': [
        'uf_ActiveAccountList', 'uf_POAccountList', 'uf_POAccountsUsed',
    ],
    'Requisition Functions': [
        'uf_FA_Requisitions', 'uf_RequisitionData', 'uf_RequisitionStatus',
        'uf_IsRequisitionLocked', 'uf_RequisitionIsVisible',
    ],
    'Budget & Reporting': [
        'uf_OrderOrBudgetBook', 'uf_UserTreeBudget', 'uf_PackCodeExport',
    ],
}


def get_format_ids_for_family(family_name):
    ids = set()
    for d in districts_by_family.get(family_name, []):
        if d.get('AccountingFormatId'):
            ids.add(d['AccountingFormatId'])
    return ids


def get_format_descs_for_family(family_name):
    descs = set()
    for d in districts_by_family.get(family_name, []):
        if d.get('AccountingSystem'):
            descs.add(d['AccountingSystem'])
    return descs


def get_req_po_for_family(family_name):
    total_reqs = 0
    total_pos = 0
    earliest = None
    latest = None
    seen = set()
    for d in districts_by_family.get(family_name, []):
        sys_name = d.get('AccountingSystem')
        if sys_name in seen:
            continue
        seen.add(sys_name)
        if sys_name in req_volumes:
            rv = req_volumes[sys_name]
            total_reqs += rv['ReqCount']
            e = rv.get('EarliestReq')
            l = rv.get('LatestReq')
            if e and (earliest is None or e < earliest):
                earliest = e
            if l and (latest is None or l > latest):
                latest = l
        if sys_name in po_volumes:
            total_pos += po_volumes[sys_name]['POCount']
    return total_reqs, total_pos, earliest, latest


# ============================================================
# GENERATE INDIVIDUAL SYSTEM DOCUMENTS (ENHANCED)
# ============================================================

def generate_system_doc(family_name, districts):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    format_ids = get_format_ids_for_family(family_name)
    format_descs = get_format_descs_for_family(family_name)
    safe_name = family_name.replace('/', '-').replace(' ', '_')

    add_title_page(doc, f'{family_name} Accounting System', 'EDS Integration Documentation')

    # TABLE OF CONTENTS
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. System Overview',
        '2. Configuration Details',
        '3. District Adoption',
        '4. Account Code Formats & Patterns',
        '5. Budget Management',
        '6. Transaction Volume & Yearly Trends',
        '7. Requisition Status Flow',
        '8. Approval Workflow',
        '9. PO Layout & Export Configuration',
        '10. Custom User Fields',
        '11. Stored Procedures & Integration Logic',
        '12. Database Functions',
        '13. Integration Notes',
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    doc.add_page_break()

    total_districts = len(districts)
    nj_count = sum(1 for d in districts if d.get('State') == 'NJ')
    ny_count = sum(1 for d in districts if d.get('State') == 'NY')
    total_schools = sum(d.get('SchoolCount', 0) for d in districts)
    total_accts = sum(d.get('AccountCount', 0) for d in districts)

    # ---- 1. SYSTEM OVERVIEW ----
    doc.add_heading('1. System Overview', level=1)

    doc.add_heading('System Family', level=2)
    doc.add_paragraph(
        f'The {family_name} accounting system is integrated with the EDS (Educational Data Services) '
        f'platform to facilitate procurement and purchase order management for school districts and '
        f'municipalities across New Jersey and New York. This integration enables requisition creation, '
        f'budget checking, approval workflows, and automated PO export to the district\'s financial system.')

    doc.add_heading('Format Variants', level=2)
    if format_descs:
        for desc in sorted(format_descs):
            doc.add_paragraph(desc, style='List Bullet')
    else:
        doc.add_paragraph('No specific format variants defined.')

    doc.add_heading('Adoption Statistics', level=2)
    stats = [
        ('Total Active Districts', str(total_districts)),
        ('New Jersey Districts', str(nj_count)),
        ('New York Districts', str(ny_count)),
        ('Other States', str(total_districts - nj_count - ny_count)),
        ('Total Schools Served', f'{total_schools:,}'),
        ('Total Active Account Codes', f'{total_accts:,}'),
    ]
    add_styled_table(doc, ['Metric', 'Value'], stats)
    doc.add_page_break()

    # ---- 2. CONFIGURATION DETAILS ----
    doc.add_heading('2. Configuration Details', level=1)
    for fid in sorted(format_ids):
        fmt = fmt_by_id.get(fid)
        if not fmt:
            continue
        doc.add_heading(f'Format: {fmt["Description"]}', level=2)
        config_rows = [
            ('Format ID', str(fid)),
            ('Short Name', str(fmt.get('ShortName') or 'N/A')),
            ('File Layout ID', str(fmt.get('FileLayoutId') or 'N/A')),
            ('Max PO Detail Items', f"{fmt.get('MaxPODetailItems'):,}" if fmt.get('MaxPODetailItems') else 'N/A'),
            ('Location Code Required', 'Yes' if fmt.get('LocationCodeRequired') else 'No'),
            ('Vendor Bid Number Required', 'Yes' if fmt.get('VendorBidNumberRequired') else 'No'),
            ('Vendor Bid Comments Required', 'Yes' if fmt.get('VendorBidCommentsRequired') else 'No'),
            ('District Accounting Code Required', 'Yes' if fmt.get('UsersDistrictAccountingCodeRequired') else 'No'),
            ('Incidental Orders Supported', 'Yes' if fmt.get('IncidentalOrdersSupported') else 'No'),
            ('Detailed Format', 'Yes' if fmt.get('DetailedFormat') else 'No (Header Only)'),
            ('Script URL', str(fmt.get('ScriptURL') or 'N/A')),
        ]
        add_styled_table(doc, ['Setting', 'Value'], config_rows)
        doc.add_paragraph()
    doc.add_page_break()

    # ---- 3. DISTRICT ADOPTION ----
    doc.add_heading('3. District Adoption', level=1)
    doc.add_heading('Complete District List', level=2)
    district_rows = []
    for d in sorted(districts, key=lambda x: x.get('Name', '')):
        district_rows.append((
            d.get('Name', ''),
            d.get('City', ''),
            d.get('State', ''),
            d.get('County', '') or '',
            str(d.get('SchoolCount', 0)),
            str(d.get('AccountCount', 0)),
            'Yes' if d.get('RequireAccounts') else 'No',
        ))
    if district_rows:
        add_styled_table(doc,
            ['District Name', 'City', 'State', 'County', 'Schools', 'Accounts', 'Req. Accts'],
            district_rows)
    else:
        doc.add_paragraph('No active districts currently using this system.')

    # County distribution
    county_counts = defaultdict(int)
    for d in districts:
        county = d.get('County') or 'Unknown'
        county_counts[county] += 1
    if len(county_counts) > 1:
        doc.add_heading('Distribution by County', level=2)
        county_rows = [(c, str(n)) for c, n in sorted(county_counts.items(), key=lambda x: -x[1])[:15]]
        add_styled_table(doc, ['County', 'Districts'], county_rows)

    doc.add_page_break()

    # ---- 4. ACCOUNT CODE FORMATS & PATTERNS ----
    doc.add_heading('4. Account Code Formats & Patterns', level=1)

    doc.add_heading('Account Code Structure', level=2)
    doc.add_paragraph(
        f'Account codes in the {family_name} system follow specific formatting conventions. '
        f'The analysis below shows the code lengths, separators, and volume of codes used.')

    # Code pattern analysis from deep data
    family_patterns = []
    for desc in format_descs:
        family_patterns.extend(code_patterns_by_system.get(desc, []))

    if family_patterns:
        doc.add_heading('Code Length Distribution', level=2)
        # Aggregate by code length
        length_counts = defaultdict(int)
        sep_set = set()
        for p in family_patterns:
            length_counts[p.get('CodeLength', 0)] += p.get('CodeCount', 0)
            if p.get('AccountSeparator'):
                sep_set.add(p['AccountSeparator'])
        length_rows = [(str(l), f'{c:,}') for l, c in sorted(length_counts.items())]
        add_styled_table(doc, ['Code Length (chars)', 'Number of Codes'], length_rows)

        if sep_set:
            doc.add_heading('Separators Used', level=2)
            for sep in sorted(sep_set):
                name = {'-': 'Dash (-)', '.': 'Period (.)', '/': 'Forward Slash (/)', ' ': 'Space'}.get(sep, f'"{sep}"')
                doc.add_paragraph(name, style='List Bullet')
    else:
        doc.add_paragraph('No account code pattern data available for this system.')

    # Show samples from top districts
    sample_districts = sorted(districts, key=lambda x: x.get('AccountCount', 0), reverse=True)[:5]
    for sd in sample_districts:
        did = sd['DistrictId']
        accts = accounts_by_district.get(did, [])[:10]
        if not accts:
            continue
        doc.add_heading(f'{sd["Name"]} ({len(accounts_by_district.get(did, []))} accounts)', level=3)
        acct_rows = [(a['Code'], a.get('AcctDesc', '')) for a in accts]
        add_styled_table(doc, ['Account Code', 'Description'], acct_rows)
        if len(accounts_by_district.get(did, [])) > 10:
            doc.add_paragraph(f'... and {len(accounts_by_district[did]) - 10} more accounts')
        doc.add_paragraph()

    doc.add_page_break()

    # ---- 5. BUDGET MANAGEMENT ----
    doc.add_heading('5. Budget Management', level=1)
    budget_districts = [d for d in districts if d['DistrictId'] in budgets_by_district]
    doc.add_paragraph(f'{len(budget_districts)} of {total_districts} districts have budget records configured.')

    if budget_districts:
        doc.add_heading('Budget Overview by District', level=2)
        budget_summary_rows = []
        for bd in sorted(budget_districts, key=lambda x: x['Name'])[:25]:
            blist = budgets_by_district[bd['DistrictId']]
            active_budgets = [b for b in blist if b.get('Active')]
            latest = max(blist, key=lambda b: b.get('StartDate') or datetime.min) if blist else None
            total_budget = sum(float(b.get('TotalBudget', 0) or 0) for b in active_budgets)
            budget_summary_rows.append((
                bd['Name'],
                str(len(blist)),
                str(len(active_budgets)),
                str(latest['StartDate'].strftime('%Y-%m-%d') if latest and latest.get('StartDate') else 'N/A'),
                f'${total_budget:,.0f}' if total_budget > 0 else 'N/A',
            ))
        add_styled_table(doc,
            ['District', 'Total Budgets', 'Active', 'Latest Start', 'Active Budget Total'],
            budget_summary_rows)

    doc.add_page_break()

    # ---- 6. TRANSACTION VOLUME & YEARLY TRENDS ----
    doc.add_heading('6. Transaction Volume & Yearly Trends', level=1)

    total_reqs, total_pos, earliest, latest = get_req_po_for_family(family_name)
    volume_rows = [
        ('Total Requisitions', f'{total_reqs:,}'),
        ('Total Purchase Orders', f'{total_pos:,}'),
        ('Earliest Transaction', str(earliest.strftime('%Y-%m-%d') if earliest else 'N/A')),
        ('Most Recent Transaction', str(latest.strftime('%Y-%m-%d') if latest else 'N/A')),
        ('Avg Requisitions per District', f'{total_reqs // max(total_districts, 1):,}'),
        ('Avg POs per District', f'{total_pos // max(total_districts, 1):,}'),
    ]
    add_styled_table(doc, ['Metric', 'Value'], volume_rows)

    # Yearly trends from deep data
    family_yearly_reqs = []
    family_yearly_pos = []
    for desc in format_descs:
        family_yearly_reqs.extend(yearly_reqs_by_system.get(desc, []))
        family_yearly_pos.extend(yearly_pos_by_system.get(desc, []))

    if family_yearly_reqs:
        doc.add_heading('Requisition Volume by Year', level=2)
        # Aggregate by year
        year_req_totals = defaultdict(lambda: {'count': 0, 'value': 0.0})
        for entry in family_yearly_reqs:
            yr = entry.get('Yr')
            if yr:
                year_req_totals[yr]['count'] += entry.get('ReqCount', 0)
                year_req_totals[yr]['value'] += float(entry.get('TotalValue', 0) or 0)

        trend_rows = []
        for yr in sorted(year_req_totals.keys()):
            data = year_req_totals[yr]
            trend_rows.append((
                str(yr),
                f'{data["count"]:,}',
                f'${data["value"]:,.0f}' if data['value'] > 0 else 'N/A',
            ))
        # Show last 15 years max
        if len(trend_rows) > 15:
            trend_rows = trend_rows[-15:]
        add_styled_table(doc, ['Year', 'Requisitions', 'Total Value'], trend_rows)

    if family_yearly_pos:
        doc.add_heading('Purchase Order Volume by Year', level=2)
        year_po_totals = defaultdict(lambda: {'count': 0, 'amount': 0.0})
        for entry in family_yearly_pos:
            yr = entry.get('Yr')
            if yr:
                year_po_totals[yr]['count'] += entry.get('POCount', 0)
                year_po_totals[yr]['amount'] += float(entry.get('TotalAmount', 0) or 0)
        po_trend_rows = []
        for yr in sorted(year_po_totals.keys()):
            data = year_po_totals[yr]
            po_trend_rows.append((
                str(yr),
                f'{data["count"]:,}',
                f'${data["amount"]:,.0f}' if data['amount'] > 0 else 'N/A',
            ))
        if len(po_trend_rows) > 15:
            po_trend_rows = po_trend_rows[-15:]
        add_styled_table(doc, ['Year', 'Purchase Orders', 'Total Amount'], po_trend_rows)

    if format_descs and len(format_descs) > 1:
        doc.add_heading('Volume by Format Variant', level=2)
        variant_rows = []
        for desc in sorted(format_descs):
            rv = req_volumes.get(desc, {})
            pv = po_volumes.get(desc, {})
            variant_rows.append((desc, f'{rv.get("ReqCount", 0):,}', f'{pv.get("POCount", 0):,}'))
        add_styled_table(doc, ['Format Variant', 'Requisitions', 'Purchase Orders'], variant_rows)

    doc.add_page_break()

    # ---- 7. REQUISITION STATUS FLOW ----
    doc.add_heading('7. Requisition Status Flow', level=1)

    doc.add_heading('Status Reference', level=2)
    doc.add_paragraph('The following statuses are defined in the EDS system for requisition tracking:')

    # Show key statuses
    key_statuses = [s for s in all_statuses if s.get('StatusId') and s.get('StatusId') <= 10]
    if key_statuses:
        status_rows = [(str(s.get('StatusId')), s.get('Name', ''), s.get('StatusCode', ''))
                       for s in key_statuses]
        add_styled_table(doc, ['ID', 'Status Name', 'Code'], status_rows)

    doc.add_heading('Typical Requisition Flow', level=2)
    flow_steps = [
        'On Hold (H) - Requisition created but not submitted',
        'Pending Approval (P) - Submitted for approval',
        'Approved (A) - All required approvals obtained',
        'At EDS (E) - Sent to EDS for PO processing',
        'PO Printed (PP) - Purchase order generated and printed',
        'Rejected (R) - Rejected by an approver (returned to requester)',
    ]
    for step in flow_steps:
        doc.add_paragraph(step, style='List Bullet')

    # Status distribution for this system
    family_status = []
    for desc in format_descs:
        family_status.extend(status_by_system.get(desc, []))

    if family_status:
        doc.add_heading('Current Status Distribution', level=2)
        # Aggregate
        status_totals = defaultdict(int)
        for entry in family_status:
            status_name = entry.get('StatusDesc', f"Status {entry.get('StatusId', '?')}")
            status_totals[status_name] += entry.get('ReqCount', 0)
        status_dist_rows = [(name, f'{count:,}') for name, count in
                            sorted(status_totals.items(), key=lambda x: -x[1])[:20]]
        add_styled_table(doc, ['Status', 'Requisition Count'], status_dist_rows)

    doc.add_page_break()

    # ---- 8. APPROVAL WORKFLOW ----
    doc.add_heading('8. Approval Workflow', level=1)

    doc.add_heading('Approval Levels', level=2)
    doc.add_paragraph(
        'EDS supports multi-level approval workflows. Each district configures its required '
        'approval level, determining how many approvals are needed before a requisition can '
        'be processed into a purchase order.')

    approval_levels = deep.get('approval_levels', [])
    if approval_levels:
        al_rows = [(str(al.get('ApprovalLevelId', '')),
                     al.get('Description', al.get('Name', 'N/A')))
                    for al in approval_levels]
        add_styled_table(doc, ['Level ID', 'Description'], al_rows)

    # District approval level distribution for this family
    family_approval = []
    for desc in format_descs:
        family_approval.extend(approval_by_system.get(desc, []))

    if family_approval:
        doc.add_heading('Approval Level Distribution', level=2)
        approval_rows = [(str(a.get('RequiredApprovalLevel', 'N/A')), str(a.get('DistrictCount', 0)))
                         for a in family_approval]
        add_styled_table(doc, ['Required Approval Level', 'Number of Districts'], approval_rows)

    doc.add_page_break()

    # ---- 9. PO LAYOUT & EXPORT ----
    doc.add_heading('9. PO Layout & Export Configuration', level=1)

    doc.add_heading('PO Form Layouts', level=2)
    doc.add_paragraph(
        'Each district has a customized PO form layout that defines the printed purchase order format. '
        'Layouts specify form dimensions, field positions, and whether continuous feed or cut sheets are used.')

    # PO layouts assigned to this family
    family_po_layouts = []
    for desc in format_descs:
        family_po_layouts.extend(po_layout_by_system.get(desc, []))

    if family_po_layouts:
        layout_rows = [(l.get('LayoutName', 'N/A'), str(l.get('DistrictCount', 0)))
                       for l in sorted(family_po_layouts, key=lambda x: -(x.get('DistrictCount') or 0))[:20]]
        add_styled_table(doc, ['PO Layout Name', 'Districts Using'], layout_rows)

    # PO Layout fields reference
    po_fields = deep.get('po_layout_fields', [])
    if po_fields:
        doc.add_heading('PO Layout Field Reference', level=2)
        doc.add_paragraph('Available fields for PO form layout customization:')
        field_rows = [(f.get('POLayoutField', ''), f.get('POLayoutSource', ''),
                       'Detail' if f.get('DetailField') else 'Header')
                      for f in po_fields[:30]]
        add_styled_table(doc, ['Field Name', 'Data Source', 'Type'], field_rows)

    doc.add_heading('Export Process', level=2)
    detailed_count = sum(1 for d in districts if d.get('DetailedFormat'))
    header_count = total_districts - detailed_count
    doc.add_paragraph(f'Districts using detailed (line-item) export: {detailed_count}')
    doc.add_paragraph(f'Districts using header-only export: {header_count}')
    doc.add_paragraph(
        f'The PO export process is handled by the usp_POPrintExport stored procedure, '
        f'which generates the accounting file for import into the {family_name} system.')

    doc.add_page_break()

    # ---- 10. CUSTOM USER FIELDS ----
    doc.add_heading('10. Custom User Fields', level=1)

    has_fields = False
    for fid in format_ids:
        fields = user_fields_by_format.get(fid, [])
        if fields:
            has_fields = True
            fmt = fmt_by_id.get(fid, {})
            doc.add_heading(f'Format: {fmt.get("Description", fid)}', level=2)
            by_district = defaultdict(list)
            for f in fields:
                by_district[f['DistrictName']].append(f)
            for dname, dfields in sorted(by_district.items()):
                doc.add_heading(dname, level=3)
                field_rows = [(str(f['FieldPos']), f['FieldName'], 'Yes' if f['RequiredField'] else 'No')
                              for f in sorted(dfields, key=lambda x: x['FieldPos'])]
                add_styled_table(doc, ['Position', 'Field Name', 'Required'], field_rows)
                doc.add_paragraph()
    if not has_fields:
        doc.add_paragraph('No custom user fields are configured for this accounting system.')

    doc.add_page_break()

    # ---- 11. STORED PROCEDURES ----
    doc.add_heading('11. Stored Procedures & Integration Logic', level=1)
    doc.add_paragraph(
        'The following stored procedures manage the core accounting integration between EDS and '
        'district financial systems. Each procedure is documented with its purpose and source code.')

    for group_name, proc_list in PROC_GROUPS.items():
        doc.add_heading(group_name, level=2)
        for proc_name in proc_list:
            source = proc_sources.get(proc_name)
            if source:
                doc.add_heading(proc_name, level=3)
                doc.add_paragraph(f'Source code length: {len(source):,} characters')
                add_code_block(doc, source, max_lines=50)
                doc.add_paragraph()

    doc.add_page_break()

    # ---- 12. DATABASE FUNCTIONS ----
    doc.add_heading('12. Database Functions', level=1)
    doc.add_paragraph(
        'These table-valued and scalar functions support requisition queries, account lookups, '
        'and budget calculations throughout the EDS application.')

    for group_name, func_list in FUNC_GROUPS.items():
        doc.add_heading(group_name, level=2)
        for fn_name in func_list:
            source = func_sources.get(fn_name)
            if source:
                doc.add_heading(fn_name, level=3)
                doc.add_paragraph(f'Source code length: {len(source):,} characters')
                add_code_block(doc, source, max_lines=40)
                doc.add_paragraph()

    doc.add_page_break()

    # ---- 13. INTEGRATION NOTES ----
    doc.add_heading('13. Integration Notes', level=1)

    doc.add_heading('Data Exchange Architecture', level=2)
    doc.add_paragraph(
        f'The {family_name} integration follows the standard EDS accounting export pipeline:\n'
        f'1. Requisitions are created and routed through the approval workflow\n'
        f'2. Approved requisitions are compiled into purchase orders\n'
        f'3. POs are formatted according to the district\'s PO layout template\n'
        f'4. The accounting export file is generated using usp_POPrintExport\n'
        f'5. The export file is transferred to the district for import into {family_name}')

    doc.add_heading('Database Tables', level=2)
    tables = [
        ('AccountingFormats', 'System format definitions and configuration flags'),
        ('AccountingDetail', 'Per-requisition line-item accounting allocations'),
        ('AccountingUserFields', 'Custom fields per district/format combination'),
        ('Accounts', 'District account codes with descriptions'),
        ('AccountSeparators', 'Valid separator characters for account codes'),
        ('BudgetAccounts', 'Budget allocations linked to accounts'),
        ('Budgets', 'Fiscal year budget definitions per district'),
        ('Ledger', 'Financial transaction journal'),
        ('Requisitions', 'Purchase requisitions linked to schools/districts'),
        ('PO', 'Purchase orders generated from requisitions'),
        ('POLayouts', 'PO form layout definitions per district'),
        ('POLayoutFields', 'Available field definitions for PO forms'),
        ('StatusTable', 'Requisition status definitions'),
        ('ApprovalLevels', 'Multi-level approval hierarchy definitions'),
        ('TransactionTypes', 'Financial transaction type classifications'),
    ]
    add_styled_table(doc, ['Table', 'Purpose'], tables)

    doc.add_heading('Charge Types', level=2)
    charge_types = deep.get('charge_types', [])
    if charge_types:
        charge_rows = [(c.get('ChargeType', ''), f'{c.get("Count", 0):,}') for c in charge_types]
        add_styled_table(doc, ['Charge Type', 'Total Records'], charge_rows)

    # Save
    filepath = os.path.join(OUTPUT_DIR, f'EDS_Accounting_{safe_name}.docx')
    doc.save(filepath)
    print(f"  Saved: {filepath}")
    return filepath


# ============================================================
# GENERATE MASTER SUMMARY (ENHANCED)
# ============================================================

def generate_master_summary():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    add_title_page(doc, 'EDS Accounting Systems', 'Master Reference Guide')

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        f'The EDS platform integrates with {len([f for f in SYSTEM_FAMILIES if f != "No Interface"])} '
        f'distinct accounting system families used by {len(all_districts)} active school districts '
        f'and municipalities across New Jersey ({sum(1 for d in all_districts if d.get("State") == "NJ")}) '
        f'and New York ({sum(1 for d in all_districts if d.get("State") == "NY")}). '
        f'These integrations enable seamless requisition-to-purchase-order workflows, '
        f'budget management, and financial data exchange.')

    doc.add_paragraph(
        f'The system manages {len(all_accounts):,} active account codes across all districts, '
        f'with {len(all_budgets):,} budget records. The platform tracks {len(all_statuses)} '
        f'distinct requisition statuses and supports {len(deep.get("approval_levels", []))} '
        f'approval levels for multi-tier authorization workflows.')

    doc.add_page_break()

    # System Comparison
    doc.add_heading('Accounting System Comparison', level=1)
    comparison_rows = []
    for family in sorted(districts_by_family.keys()):
        dists = districts_by_family[family]
        total_reqs, total_pos, earliest, latest = get_req_po_for_family(family)
        comparison_rows.append((
            family, str(len(dists)),
            f'{sum(d.get("AccountCount", 0) for d in dists):,}',
            f'{total_reqs:,}', f'{total_pos:,}',
        ))
    comparison_rows.sort(key=lambda x: int(x[1]), reverse=True)
    add_styled_table(doc, ['System', 'Districts', 'Accounts', 'Requisitions', 'POs'], comparison_rows)

    doc.add_page_break()

    # Market Share
    doc.add_heading('Market Share Analysis', level=1)
    doc.add_heading('By District Count', level=2)
    total = len(all_districts)
    market_rows = []
    for family in sorted(districts_by_family.keys(), key=lambda f: len(districts_by_family[f]), reverse=True):
        count = len(districts_by_family[family])
        pct = count / total * 100
        market_rows.append((family, str(count), f'{pct:.1f}%'))
    add_styled_table(doc, ['System', 'Districts', 'Market Share'], market_rows)

    doc.add_heading('Geographic Distribution', level=2)
    geo_rows = []
    for family in sorted(districts_by_family.keys(), key=lambda f: len(districts_by_family[f]), reverse=True):
        dists = districts_by_family[family]
        nj = sum(1 for d in dists if d.get('State') == 'NJ')
        ny = sum(1 for d in dists if d.get('State') == 'NY')
        other = len(dists) - nj - ny
        geo_rows.append((family, str(nj), str(ny), str(other)))
    add_styled_table(doc, ['System', 'NJ', 'NY', 'Other'], geo_rows)

    doc.add_page_break()

    # Approval Workflow Overview
    doc.add_heading('Approval Workflow Overview', level=1)
    doc.add_paragraph(
        'The EDS platform supports configurable multi-level approval workflows. '
        'Districts set their required approval level, and requisitions must pass through '
        'each level before being eligible for PO generation.')

    approval_levels = deep.get('approval_levels', [])
    if approval_levels:
        al_rows = [(str(al.get('ApprovalLevelId', '')),
                     al.get('Description', al.get('Name', 'N/A')))
                    for al in approval_levels]
        add_styled_table(doc, ['Level', 'Description'], al_rows)

    doc.add_heading('Complete Status Reference', level=2)
    if all_statuses:
        status_rows = [(str(s.get('StatusId', '')), s.get('Name', ''), s.get('StatusCode', ''))
                       for s in all_statuses if s.get('Name')]
        add_styled_table(doc, ['ID', 'Status', 'Code'], status_rows)

    doc.add_page_break()

    # Transaction Types & Charges
    doc.add_heading('Transaction Types & District Charges', level=1)

    doc.add_heading('Transaction Types', level=2)
    tx_types = deep.get('transaction_types', [])
    if tx_types:
        tx_rows = [(str(t.get('TransactionTypeId', '')), t.get('Description', ''))
                    for t in tx_types]
        add_styled_table(doc, ['Type ID', 'Description'], tx_rows)

    doc.add_heading('District Charge Types', level=2)
    charge_types = deep.get('charge_types', [])
    if charge_types:
        charge_rows = [(c.get('ChargeType', ''), f'{c.get("Count", 0):,}') for c in charge_types]
        add_styled_table(doc, ['Charge Type', 'Total Records'], charge_rows)

    doc.add_page_break()

    # Format Configuration Matrix
    doc.add_heading('Format Configuration Matrix', level=1)
    doc.add_paragraph('Feature requirements for each active format variant:')
    active_format_ids = set()
    for d in all_districts:
        if d.get('AccountingFormatId'):
            active_format_ids.add(d['AccountingFormatId'])

    config_rows = []
    for fid in sorted(active_format_ids):
        fmt = fmt_by_id.get(fid)
        if not fmt:
            continue
        district_count = sum(1 for d in all_districts if d.get('AccountingFormatId') == fid)
        if district_count == 0:
            continue
        y_n = lambda v: 'Y' if v else 'N'
        config_rows.append((
            fmt['Description'][:40], str(district_count),
            y_n(fmt.get('LocationCodeRequired')), y_n(fmt.get('VendorBidNumberRequired')),
            y_n(fmt.get('IncidentalOrdersSupported')), y_n(fmt.get('DetailedFormat')),
        ))
    config_rows.sort(key=lambda x: int(x[1]), reverse=True)
    add_styled_table(doc, ['Format', '#Dist', 'LocCode', 'BidNum', 'Incidental', 'Detailed'], config_rows)

    doc.add_page_break()

    # PO Layout Summary
    doc.add_heading('PO Layout Summary', level=1)
    po_layouts = deep.get('po_layouts', [])
    doc.add_paragraph(f'Total PO layouts defined: {len(po_layouts)}')
    doc.add_paragraph(
        'Each district has a customized PO form layout that defines paper size, orientation, '
        'field positions, and whether continuous feed or laser printing is used.')

    po_fields = deep.get('po_layout_fields', [])
    if po_fields:
        doc.add_heading('Available PO Fields', level=2)
        field_rows = [(f.get('POLayoutField', ''), f.get('POLayoutSource', ''),
                       'Detail' if f.get('DetailField') else 'Header')
                      for f in po_fields]
        add_styled_table(doc, ['Field Name', 'Data Source', 'Type'], field_rows)

    doc.add_page_break()

    # Stored Procedure Reference
    doc.add_heading('Stored Procedure Reference', level=1)
    doc.add_paragraph(
        f'The EDS system uses {len(proc_sources)} stored procedures and {len(func_sources)} '
        f'functions for accounting integration.')

    for group_name, proc_list in PROC_GROUPS.items():
        doc.add_heading(group_name, level=2)
        for proc_name in proc_list:
            source = proc_sources.get(proc_name)
            if source:
                doc.add_paragraph(f'{proc_name} ({len(source):,} chars)', style='List Bullet')

    doc.add_heading('Database Functions', level=2)
    for group_name, func_list in FUNC_GROUPS.items():
        doc.add_heading(group_name, level=3)
        for fn_name in func_list:
            source = func_sources.get(fn_name)
            if source:
                doc.add_paragraph(f'{fn_name} ({len(source):,} chars)', style='List Bullet')

    doc.add_page_break()

    # Database Schema Reference
    doc.add_heading('Database Schema Reference', level=1)
    doc.add_heading('Core Accounting Tables', level=2)
    schema_rows = [
        ('AccountingFormats', '49', 'System format definitions and integration settings'),
        ('AccountingDetail', '-', 'Per-requisition line-item accounting allocations'),
        ('AccountingUserFields', f'{len(all_user_fields)}', 'Custom fields per district/format'),
        ('Accounts', f'{len(all_accounts):,}', 'District account codes with descriptions'),
        ('BudgetAccounts', '-', 'Budget allocations linked to accounts'),
        ('Budgets', f'{len(all_budgets):,}', 'Fiscal year budget definitions per district'),
        ('Ledger', '-', 'Financial transaction journal'),
        ('POLayouts', f'{len(po_layouts)}', 'PO form layout templates'),
        ('POLayoutFields', f'{len(po_fields)}', 'Available PO form fields'),
        ('StatusTable', f'{len(all_statuses)}', 'Requisition status definitions'),
        ('ApprovalLevels', f'{len(approval_levels)}', 'Approval hierarchy levels'),
    ]
    add_styled_table(doc, ['Table', 'Records', 'Description'], schema_rows)

    doc.add_heading('Related Tables', level=2)
    related = [
        ('District', f'{len(all_districts)}', 'District master - links to AccountingFormats'),
        ('Requisitions', '-', 'Purchase requisitions with status tracking'),
        ('PO', '-', 'Purchase orders from approved requisitions'),
        ('Detail', '-', 'Requisition line items with pricing'),
        ('School', '-', 'Schools within districts'),
        ('TransactionTypes', f'{len(tx_types)}', 'Financial transaction classifications'),
    ]
    add_styled_table(doc, ['Table', 'Active Records', 'Description'], related)

    doc.add_page_break()

    # Complete District Directory
    doc.add_heading('Complete District Directory', level=1)
    doc.add_paragraph(f'All {len(all_districts)} active districts with their assigned accounting system.')
    all_dist_rows = []
    for d in sorted(all_districts, key=lambda x: x.get('Name', '')):
        all_dist_rows.append((
            d.get('Name', ''), d.get('City', ''), d.get('State', ''),
            (d.get('AccountingSystem') or 'None')[:35], str(d.get('AccountCount', 0)),
        ))
    add_styled_table(doc, ['District', 'City', 'ST', 'Accounting System', 'Accts'], all_dist_rows)

    filepath = os.path.join(OUTPUT_DIR, 'EDS_Accounting_Master_Summary.docx')
    doc.save(filepath)
    print(f"  Saved: {filepath}")
    return filepath


# ============================================================
# MAIN EXECUTION
# ============================================================
print("\n=== Generating Enhanced Individual System Documents ===")
generated = []

for family in sorted(districts_by_family.keys()):
    districts = districts_by_family[family]
    print(f"\nGenerating: {family} ({len(districts)} districts)...")
    try:
        path = generate_system_doc(family, districts)
        generated.append(path)
    except Exception as e:
        print(f"  ERROR: {e}")
        import traceback
        traceback.print_exc()

print("\n=== Generating Enhanced Master Summary ===")
try:
    master_path = generate_master_summary()
    generated.append(master_path)
except Exception as e:
    print(f"  ERROR: {e}")
    import traceback
    traceback.print_exc()

conn.close()

print(f"\n{'='*60}")
print(f"Enhanced documentation generation complete!")
print(f"Total documents generated: {len(generated)}")
print(f"Output directory: {os.path.abspath(OUTPUT_DIR)}")
for g in generated:
    size = os.path.getsize(g)
    print(f"  {os.path.basename(g)} ({size:,} bytes)")
