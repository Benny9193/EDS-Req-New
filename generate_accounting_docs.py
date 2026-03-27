"""Generate comprehensive DOCX documentation for each EDS accounting system."""
import pyodbc
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from collections import defaultdict
from datetime import datetime
import decimal

# Database connection
conn_str = (
    'DRIVER={ODBC Driver 17 for SQL Server};'
    'SERVER=eds-sqlserver.eastus2.cloudapp.azure.com;'
    'DATABASE=EDS;'
    'UID=EDSAdmin;'
    'PWD=Consultant~!;'
    'TrustServerCertificate=yes;'
)

OUTPUT_DIR = 'docs/accounting'
os.makedirs(OUTPUT_DIR, exist_ok=True)

conn = pyodbc.connect(conn_str)
cursor = conn.cursor()


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    from docx.oxml.ns import qn
    from lxml import etree
    shading = etree.SubElement(cell._tc.get_or_add_tcPr(), qn('w:shd'))
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a formatted table to a document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '1c1a83')  # EDS primary blue

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, value in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(value) if value is not None else ''
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'f0f0f8')

    return table


def add_title_page(doc, title, subtitle=''):
    """Add a title page to the document."""
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(28, 26, 131)  # EDS primary

    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(subtitle)
        run2.font.size = Pt(14)
        run2.font.color.rgb = RGBColor(74, 72, 144)  # EDS secondary

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run('Educational Data Services (EDS)')
    run3.font.size = Pt(12)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
    run4.font.size = Pt(10)
    run4.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()


# ============================================================
# GATHER ALL DATA
# ============================================================
print("Gathering data from database...")

# Accounting formats
cursor.execute("SELECT * FROM dbo.AccountingFormats ORDER BY AccountingFormatId")
fmt_cols = [d[0] for d in cursor.description]
all_formats = [dict(zip(fmt_cols, row)) for row in cursor.fetchall()]
fmt_by_id = {f['AccountingFormatId']: f for f in all_formats}

# Districts with their accounting system
cursor.execute("""
    SELECT d.DistrictId, d.DistrictCode, d.Name, d.City, d.State, d.County,
           d.AccountingFormatId, d.RequireAccounts, d.CurrentBudgetOnly,
           d.AccountSeparator, d.AccountingDistrictCode, d.AccountingSystemOptions,
           af.Description AS AccountingSystem, af.ShortName,
           af.LocationCodeRequired, af.VendorBidNumberRequired,
           af.VendorBidCommentsRequired, af.UsersDistrictAccountingCodeRequired,
           af.IncidentalOrdersSupported, af.DetailedFormat,
           (SELECT COUNT(*) FROM dbo.School s WHERE s.DistrictId = d.DistrictId) AS SchoolCount,
           (SELECT COUNT(*) FROM dbo.Accounts a WHERE a.DistrictId = d.DistrictId AND a.Active = 1) AS AccountCount
    FROM dbo.District d
    LEFT JOIN dbo.AccountingFormats af ON d.AccountingFormatId = af.AccountingFormatId
    WHERE d.Active = 1
    ORDER BY af.Description, d.Name
""")
dist_cols = [d[0] for d in cursor.description]
all_districts = [dict(zip(dist_cols, row)) for row in cursor.fetchall()]

# Group districts by system family
SYSTEM_FAMILIES = {
    'Finance Manager': lambda d: 'Finance Manager' in (d.get('AccountingSystem') or ''),
    'Systems 3000': lambda d: 'System' in (d.get('AccountingSystem') or '') and '3000' in (d.get('AccountingSystem') or ''),
    'CSI': lambda d: 'CSI' in (d.get('AccountingSystem') or ''),
    'Genesis': lambda d: 'Genesis' in (d.get('AccountingSystem') or ''),
    'WinCap': lambda d: 'Wincap' in (d.get('AccountingSystem') or '') or 'WinCap' in (d.get('ShortName') or ''),
    'CDK': lambda d: 'CDK' in (d.get('AccountingSystem') or ''),
    'Edu-Met': lambda d: 'Edu-Met' in (d.get('AccountingSystem') or '') or 'EMapp' in (d.get('AccountingSystem') or ''),
    'Sage/Alio': lambda d: 'Sage' in (d.get('AccountingSystem') or '') or 'Alio' in (d.get('AccountingSystem') or ''),
    'Edmunds': lambda d: 'Edmunds' in (d.get('AccountingSystem') or ''),
    'Infinite Visions': lambda d: 'Infinite' in (d.get('AccountingSystem') or ''),
    'No Interface': lambda d: 'No Accounting' in (d.get('AccountingSystem') or ''),
}

districts_by_family = defaultdict(list)
for d in all_districts:
    matched = False
    for family, predicate in SYSTEM_FAMILIES.items():
        if predicate(d):
            districts_by_family[family].append(d)
            matched = True
            break
    if not matched:
        sys_name = d.get('AccountingSystem') or 'Unassigned'
        districts_by_family[sys_name].append(d)

# Account codes
cursor.execute("""
    SELECT d.DistrictId, af.Description AS SystemName, d.Name AS DistrictName,
           a.Code, a.Description AS AcctDesc
    FROM dbo.Accounts a
    JOIN dbo.District d ON a.DistrictId = d.DistrictId
    LEFT JOIN dbo.AccountingFormats af ON d.AccountingFormatId = af.AccountingFormatId
    WHERE a.Active = 1 AND d.Active = 1
""")
acct_cols = [d[0] for d in cursor.description]
all_accounts = [dict(zip(acct_cols, row)) for row in cursor.fetchall()]

accounts_by_district = defaultdict(list)
for a in all_accounts:
    accounts_by_district[a['DistrictId']].append(a)

# Budget data
cursor.execute("""
    SELECT d.DistrictId, d.Name AS DistrictName, af.Description AS SystemName,
           b.Name AS BudgetName, b.Active, b.StartDate, b.EndDate,
           COUNT(ba.BudgetAccountId) AS AccountCount,
           SUM(ba.BudgetAmount) AS TotalBudget
    FROM dbo.Budgets b
    JOIN dbo.District d ON b.DistrictId = d.DistrictId
    LEFT JOIN dbo.AccountingFormats af ON d.AccountingFormatId = af.AccountingFormatId
    LEFT JOIN dbo.BudgetAccounts ba ON b.BudgetId = ba.BudgetId
    WHERE d.Active = 1
    GROUP BY d.DistrictId, d.Name, af.Description, b.Name, b.Active, b.StartDate, b.EndDate
    ORDER BY d.Name, b.StartDate DESC
""")
bdg_cols = [d[0] for d in cursor.description]
all_budgets = [dict(zip(bdg_cols, row)) for row in cursor.fetchall()]

budgets_by_district = defaultdict(list)
for b in all_budgets:
    budgets_by_district[b['DistrictId']].append(b)

# Requisition volumes
cursor.execute("""
    SELECT af.Description AS SystemName,
           COUNT(r.RequisitionId) AS ReqCount,
           MIN(r.DateEntered) AS EarliestReq,
           MAX(r.DateEntered) AS LatestReq
    FROM dbo.Requisitions r
    JOIN dbo.School s ON r.SchoolId = s.SchoolId
    JOIN dbo.District d ON s.DistrictId = d.DistrictId
    LEFT JOIN dbo.AccountingFormats af ON d.AccountingFormatId = af.AccountingFormatId
    WHERE d.Active = 1
    GROUP BY af.Description
""")
req_cols = [d[0] for d in cursor.description]
req_volumes = {row[0]: dict(zip(req_cols, row)) for row in cursor.fetchall()}

# PO volumes
cursor.execute("""
    SELECT af.Description AS SystemName, COUNT(p.POId) AS POCount
    FROM dbo.PO p
    JOIN dbo.Requisitions r ON p.RequisitionId = r.RequisitionId
    JOIN dbo.School s ON r.SchoolId = s.SchoolId
    JOIN dbo.District d ON s.DistrictId = d.DistrictId
    LEFT JOIN dbo.AccountingFormats af ON d.AccountingFormatId = af.AccountingFormatId
    WHERE d.Active = 1
    GROUP BY af.Description
""")
po_cols = [d[0] for d in cursor.description]
po_volumes = {row[0]: dict(zip(po_cols, row)) for row in cursor.fetchall()}

# User fields
cursor.execute("""
    SELECT auf.AccountingFormatId, af.Description AS FormatName,
           auf.DistrictId, d.Name AS DistrictName,
           auf.FieldPos, auf.FieldName, auf.RequiredField
    FROM dbo.AccountingUserFields auf
    JOIN dbo.AccountingFormats af ON auf.AccountingFormatId = af.AccountingFormatId
    JOIN dbo.District d ON auf.DistrictId = d.DistrictId
    ORDER BY auf.AccountingFormatId, auf.DistrictId, auf.FieldPos
""")
uf_cols = [d[0] for d in cursor.description]
all_user_fields = [dict(zip(uf_cols, row)) for row in cursor.fetchall()]

user_fields_by_format = defaultdict(list)
for uf in all_user_fields:
    user_fields_by_format[uf['AccountingFormatId']].append(uf)

print(f"Data gathered: {len(all_districts)} districts, {len(all_accounts)} accounts, {len(all_budgets)} budgets")


# ============================================================
# GENERATE INDIVIDUAL SYSTEM DOCUMENTS
# ============================================================

def get_format_ids_for_family(family_name):
    """Get all AccountingFormatId values that belong to a system family."""
    ids = set()
    for d in districts_by_family.get(family_name, []):
        if d.get('AccountingFormatId'):
            ids.add(d['AccountingFormatId'])
    return ids


def get_req_po_for_family(family_name):
    """Aggregate req/PO counts for a system family's format descriptions."""
    total_reqs = 0
    total_pos = 0
    earliest = None
    latest = None
    for d in districts_by_family.get(family_name, []):
        sys_name = d.get('AccountingSystem')
        if sys_name in req_volumes:
            rv = req_volumes[sys_name]
            total_reqs += rv['ReqCount']
            e = rv.get('EarliestReq')
            l = rv.get('LatestReq')
            if e and (earliest is None or e < earliest):
                earliest = e
            if l and (latest is None or l > latest):
                latest = l
        if sys_name in po_volumes:
            total_pos += po_volumes[sys_name]['POCount']
    return total_reqs, total_pos, earliest, latest


def generate_system_doc(family_name, districts):
    """Generate a comprehensive DOCX for one accounting system family."""
    doc = Document()

    # Styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    format_ids = get_format_ids_for_family(family_name)
    format_descs = set()
    for d in districts:
        if d.get('AccountingSystem'):
            format_descs.add(d['AccountingSystem'])

    safe_name = family_name.replace('/', '-').replace(' ', '_')

    # Title page
    add_title_page(doc, f'{family_name} Accounting System', 'EDS Integration Documentation')

    # ---- TABLE OF CONTENTS placeholder ----
    doc.add_heading('Table of Contents', level=1)
    doc.add_paragraph('1. System Overview')
    doc.add_paragraph('2. Configuration Details')
    doc.add_paragraph('3. District Adoption')
    doc.add_paragraph('4. Account Code Formats')
    doc.add_paragraph('5. Budget Management')
    doc.add_paragraph('6. Transaction Volume')
    doc.add_paragraph('7. Custom User Fields')
    doc.add_paragraph('8. Integration Notes')
    doc.add_page_break()

    # ---- 1. SYSTEM OVERVIEW ----
    doc.add_heading('1. System Overview', level=1)

    doc.add_heading('System Family', level=2)
    doc.add_paragraph(f'The {family_name} accounting system is integrated with the EDS (Educational Data Services) '
                      f'platform to facilitate procurement and purchase order management for school districts and '
                      f'municipalities across New Jersey and New York.')

    doc.add_heading('Format Variants', level=2)
    if format_descs:
        for desc in sorted(format_descs):
            doc.add_paragraph(desc, style='List Bullet')
    else:
        doc.add_paragraph('No specific format variants defined.')

    doc.add_heading('Adoption Statistics', level=2)
    total_districts = len(districts)
    nj_count = sum(1 for d in districts if d.get('State') == 'NJ')
    ny_count = sum(1 for d in districts if d.get('State') == 'NY')
    total_schools = sum(d.get('SchoolCount', 0) for d in districts)
    total_accts = sum(d.get('AccountCount', 0) for d in districts)

    stats = [
        ('Total Active Districts', str(total_districts)),
        ('New Jersey Districts', str(nj_count)),
        ('New York Districts', str(ny_count)),
        ('Other States', str(total_districts - nj_count - ny_count)),
        ('Total Schools Served', f'{total_schools:,}'),
        ('Total Active Account Codes', f'{total_accts:,}'),
    ]
    add_styled_table(doc, ['Metric', 'Value'], stats)

    doc.add_page_break()

    # ---- 2. CONFIGURATION DETAILS ----
    doc.add_heading('2. Configuration Details', level=1)

    for fid in sorted(format_ids):
        fmt = fmt_by_id.get(fid)
        if not fmt:
            continue

        doc.add_heading(f'Format: {fmt["Description"]}', level=2)
        config_rows = [
            ('Format ID', str(fid)),
            ('Short Name', str(fmt.get('ShortName') or 'N/A')),
            ('File Layout ID', str(fmt.get('FileLayoutId') or 'N/A')),
            ('Max PO Detail Items', f"{fmt.get('MaxPODetailItems'):,}" if fmt.get('MaxPODetailItems') else 'N/A'),
            ('Location Code Required', 'Yes' if fmt.get('LocationCodeRequired') else 'No'),
            ('Vendor Bid Number Required', 'Yes' if fmt.get('VendorBidNumberRequired') else 'No'),
            ('Vendor Bid Comments Required', 'Yes' if fmt.get('VendorBidCommentsRequired') else 'No'),
            ('District Accounting Code Required', 'Yes' if fmt.get('UsersDistrictAccountingCodeRequired') else 'No'),
            ('Incidental Orders Supported', 'Yes' if fmt.get('IncidentalOrdersSupported') else 'No'),
            ('Detailed Format', 'Yes' if fmt.get('DetailedFormat') else 'No (Header Only)'),
            ('Script URL', str(fmt.get('ScriptURL') or 'N/A')),
        ]
        add_styled_table(doc, ['Setting', 'Value'], config_rows)
        doc.add_paragraph()

    doc.add_page_break()

    # ---- 3. DISTRICT ADOPTION ----
    doc.add_heading('3. District Adoption', level=1)

    doc.add_heading('Complete District List', level=2)
    district_rows = []
    for d in sorted(districts, key=lambda x: x.get('Name', '')):
        district_rows.append((
            d.get('Name', ''),
            d.get('City', ''),
            d.get('State', ''),
            d.get('County', '') or '',
            str(d.get('SchoolCount', 0)),
            str(d.get('AccountCount', 0)),
            'Yes' if d.get('RequireAccounts') else 'No',
        ))

    if district_rows:
        add_styled_table(doc,
            ['District Name', 'City', 'State', 'County', 'Schools', 'Accounts', 'Req. Accts'],
            district_rows)
    else:
        doc.add_paragraph('No active districts currently using this system.')

    doc.add_page_break()

    # ---- 4. ACCOUNT CODE FORMATS ----
    doc.add_heading('4. Account Code Formats', level=1)

    doc.add_heading('Account Code Structure', level=2)
    doc.add_paragraph(
        f'Account codes in the {family_name} system follow specific formatting conventions. '
        f'Each district may use slight variations but generally adheres to a common pattern '
        f'dictated by the accounting system requirements.')

    # Show samples from up to 5 districts
    sample_districts = sorted(districts, key=lambda x: x.get('AccountCount', 0), reverse=True)[:5]
    for sd in sample_districts:
        did = sd['DistrictId']
        accts = accounts_by_district.get(did, [])[:10]
        if not accts:
            continue

        doc.add_heading(f'{sd["Name"]} ({len(accounts_by_district.get(did, []))} accounts)', level=3)
        acct_rows = [(a['Code'], a.get('AcctDesc', '')) for a in accts]
        add_styled_table(doc, ['Account Code', 'Description'], acct_rows)
        if len(accounts_by_district.get(did, [])) > 10:
            doc.add_paragraph(f'... and {len(accounts_by_district[did]) - 10} more accounts', style='List Bullet')
        doc.add_paragraph()

    # Separator analysis
    separators = set()
    for d in districts:
        if d.get('AccountSeparator'):
            separators.add(d['AccountSeparator'])
    if separators:
        doc.add_heading('Account Code Separators Used', level=2)
        for sep in sorted(separators):
            name = {'-': 'Dash (-)', '.': 'Period (.)', '/': 'Forward Slash (/)', ' ': 'Space'}.get(sep, f'"{sep}"')
            doc.add_paragraph(name, style='List Bullet')

    doc.add_page_break()

    # ---- 5. BUDGET MANAGEMENT ----
    doc.add_heading('5. Budget Management', level=1)

    # Show budget summary for top districts
    budget_districts = [d for d in districts if d['DistrictId'] in budgets_by_district]
    doc.add_paragraph(f'{len(budget_districts)} of {total_districts} districts have budget records configured.')

    if budget_districts:
        doc.add_heading('Budget Overview by District', level=2)
        budget_summary_rows = []
        for bd in sorted(budget_districts, key=lambda x: x['Name'])[:20]:
            blist = budgets_by_district[bd['DistrictId']]
            active_budgets = [b for b in blist if b.get('Active')]
            latest = max(blist, key=lambda b: b.get('StartDate') or datetime.min) if blist else None
            budget_summary_rows.append((
                bd['Name'],
                str(len(blist)),
                str(len(active_budgets)),
                str(latest['StartDate'].strftime('%Y-%m-%d') if latest and latest.get('StartDate') else 'N/A'),
                str(latest.get('AccountCount', 0)) if latest else '0',
            ))

        add_styled_table(doc,
            ['District', 'Total Budgets', 'Active', 'Latest Start', 'Accounts'],
            budget_summary_rows)

    doc.add_page_break()

    # ---- 6. TRANSACTION VOLUME ----
    doc.add_heading('6. Transaction Volume', level=1)

    total_reqs, total_pos, earliest, latest = get_req_po_for_family(family_name)

    volume_rows = [
        ('Total Requisitions', f'{total_reqs:,}'),
        ('Total Purchase Orders', f'{total_pos:,}'),
        ('Earliest Transaction', str(earliest.strftime('%Y-%m-%d') if earliest else 'N/A')),
        ('Most Recent Transaction', str(latest.strftime('%Y-%m-%d') if latest else 'N/A')),
        ('Avg Requisitions per District', f'{total_reqs // max(total_districts, 1):,}'),
        ('Avg POs per District', f'{total_pos // max(total_districts, 1):,}'),
    ]
    add_styled_table(doc, ['Metric', 'Value'], volume_rows)

    if format_descs:
        doc.add_heading('Volume by Format Variant', level=2)
        variant_rows = []
        for desc in sorted(format_descs):
            rv = req_volumes.get(desc, {})
            pv = po_volumes.get(desc, {})
            variant_rows.append((
                desc,
                f'{rv.get("ReqCount", 0):,}',
                f'{pv.get("POCount", 0):,}',
            ))
        add_styled_table(doc, ['Format Variant', 'Requisitions', 'Purchase Orders'], variant_rows)

    doc.add_page_break()

    # ---- 7. CUSTOM USER FIELDS ----
    doc.add_heading('7. Custom User Fields', level=1)

    has_fields = False
    for fid in format_ids:
        fields = user_fields_by_format.get(fid, [])
        if fields:
            has_fields = True
            fmt = fmt_by_id.get(fid, {})
            doc.add_heading(f'Format: {fmt.get("Description", fid)}', level=2)

            # Group by district
            by_district = defaultdict(list)
            for f in fields:
                by_district[f['DistrictName']].append(f)

            for dname, dfields in sorted(by_district.items()):
                doc.add_heading(dname, level=3)
                field_rows = [(str(f['FieldPos']), f['FieldName'], 'Yes' if f['RequiredField'] else 'No')
                              for f in sorted(dfields, key=lambda x: x['FieldPos'])]
                add_styled_table(doc, ['Position', 'Field Name', 'Required'], field_rows)
                doc.add_paragraph()

    if not has_fields:
        doc.add_paragraph('No custom user fields are configured for this accounting system.')

    doc.add_page_break()

    # ---- 8. INTEGRATION NOTES ----
    doc.add_heading('8. Integration Notes', level=1)

    doc.add_heading('Data Exchange', level=2)
    detailed_count = sum(1 for d in districts if d.get('DetailedFormat'))
    header_count = total_districts - detailed_count
    doc.add_paragraph(f'Districts using detailed (line-item) export: {detailed_count}')
    doc.add_paragraph(f'Districts using header-only export: {header_count}')

    doc.add_heading('Database Tables', level=2)
    doc.add_paragraph('The following EDS database tables are involved in the accounting integration:')
    tables = [
        ('AccountingFormats', 'System format definitions and configuration'),
        ('AccountingDetail', 'Line-item accounting detail per requisition'),
        ('AccountingUserFields', 'Custom fields per district/format'),
        ('Accounts', 'District account codes'),
        ('AccountSeparators', 'Account code separator characters'),
        ('BudgetAccounts', 'Budget allocations per account'),
        ('Budgets', 'Fiscal year budget definitions'),
        ('Ledger', 'Financial transaction log'),
        ('Requisitions', 'Purchase requisitions'),
        ('PO', 'Purchase orders generated from requisitions'),
    ]
    add_styled_table(doc, ['Table', 'Purpose'], tables)

    doc.add_heading('Key Stored Procedures', level=2)
    doc.add_paragraph('The following stored procedures manage accounting operations:')
    procs = [
        'sp_FA_AddUpdateAccountCode', 'sp_FA_AvailableAccounts',
        'sp_FA_SetBudgetAccount', 'sp_FA_SetUserAccount',
        'sp_RefreshAccounts', 'sp_MergeAccounts',
        'sp_ResetDistrictAccountingYear', 'sp_SetBudgetYear',
        'usp_BringAccountsForward', 'usp_UpdateBudgets',
        'usp_POPrintExport',
    ]
    for p in procs:
        doc.add_paragraph(p, style='List Bullet')

    # Save
    filepath = os.path.join(OUTPUT_DIR, f'EDS_Accounting_{safe_name}.docx')
    doc.save(filepath)
    print(f"  Saved: {filepath}")
    return filepath


# ============================================================
# GENERATE MASTER SUMMARY DOCUMENT
# ============================================================

def generate_master_summary():
    """Generate the master summary document covering all systems."""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    add_title_page(doc, 'EDS Accounting Systems', 'Master Reference Guide')

    # ---- Executive Summary ----
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        f'The EDS platform integrates with {len([f for f in SYSTEM_FAMILIES if f != "No Interface"])} '
        f'distinct accounting system families used by {len(all_districts)} active school districts '
        f'and municipalities across New Jersey ({sum(1 for d in all_districts if d.get("State") == "NJ")}) '
        f'and New York ({sum(1 for d in all_districts if d.get("State") == "NY")}). '
        f'These integrations enable seamless requisition-to-purchase-order workflows, '
        f'budget management, and financial data exchange between EDS and each district\'s '
        f'accounting software.')

    doc.add_paragraph(
        f'The system manages {len(all_accounts):,} active account codes across all districts, '
        f'with {len(all_budgets):,} budget records spanning multiple fiscal years.')

    doc.add_page_break()

    # ---- System Comparison ----
    doc.add_heading('Accounting System Comparison', level=1)

    comparison_rows = []
    for family in sorted(districts_by_family.keys()):
        dists = districts_by_family[family]
        total_reqs, total_pos, earliest, latest = get_req_po_for_family(family)
        comparison_rows.append((
            family,
            str(len(dists)),
            f'{sum(d.get("AccountCount", 0) for d in dists):,}',
            f'{total_reqs:,}',
            f'{total_pos:,}',
        ))

    comparison_rows.sort(key=lambda x: int(x[1]), reverse=True)
    add_styled_table(doc,
        ['System', 'Districts', 'Accounts', 'Requisitions', 'POs'],
        comparison_rows)

    doc.add_page_break()

    # ---- Market Share ----
    doc.add_heading('Market Share Analysis', level=1)

    doc.add_heading('By District Count', level=2)
    total = len(all_districts)
    market_rows = []
    for family in sorted(districts_by_family.keys(), key=lambda f: len(districts_by_family[f]), reverse=True):
        count = len(districts_by_family[family])
        pct = count / total * 100
        market_rows.append((family, str(count), f'{pct:.1f}%'))

    add_styled_table(doc, ['System', 'Districts', 'Market Share'], market_rows)

    doc.add_heading('Geographic Distribution', level=2)
    geo_rows = []
    for family in sorted(districts_by_family.keys(), key=lambda f: len(districts_by_family[f]), reverse=True):
        dists = districts_by_family[family]
        nj = sum(1 for d in dists if d.get('State') == 'NJ')
        ny = sum(1 for d in dists if d.get('State') == 'NY')
        other = len(dists) - nj - ny
        geo_rows.append((family, str(nj), str(ny), str(other)))

    add_styled_table(doc, ['System', 'NJ', 'NY', 'Other'], geo_rows)

    doc.add_page_break()

    # ---- Format Configuration Matrix ----
    doc.add_heading('Format Configuration Matrix', level=1)

    doc.add_paragraph('This matrix shows which features are required/supported for each active format variant.')

    active_format_ids = set()
    for d in all_districts:
        if d.get('AccountingFormatId'):
            active_format_ids.add(d['AccountingFormatId'])

    config_rows = []
    for fid in sorted(active_format_ids):
        fmt = fmt_by_id.get(fid)
        if not fmt:
            continue
        district_count = sum(1 for d in all_districts if d.get('AccountingFormatId') == fid)
        if district_count == 0:
            continue
        y_n = lambda v: 'Y' if v else 'N'
        config_rows.append((
            fmt['Description'][:40],
            str(district_count),
            y_n(fmt.get('LocationCodeRequired')),
            y_n(fmt.get('VendorBidNumberRequired')),
            y_n(fmt.get('IncidentalOrdersSupported')),
            y_n(fmt.get('DetailedFormat')),
        ))

    config_rows.sort(key=lambda x: int(x[1]), reverse=True)
    add_styled_table(doc,
        ['Format', '#Dist', 'LocCode', 'BidNum', 'Incidental', 'Detailed'],
        config_rows)

    doc.add_page_break()

    # ---- Database Schema ----
    doc.add_heading('Database Schema Reference', level=1)

    doc.add_heading('Core Accounting Tables', level=2)
    schema_rows = [
        ('AccountingFormats', '49', 'Accounting system definitions and integration settings'),
        ('AccountingDetail', '-', 'Per-requisition line-item accounting allocations'),
        ('AccountingUserFields', f'{len(all_user_fields)}', 'Custom fields per district/format combination'),
        ('Accounts', f'{len(all_accounts):,}', 'District account codes with descriptions'),
        ('AccountSeparators', '-', 'Valid separator characters for account codes'),
        ('BudgetAccounts', '-', 'Budget allocations linked to accounts'),
        ('Budgets', f'{len(all_budgets):,}', 'Fiscal year budget definitions per district'),
        ('Ledger', '-', 'Financial transaction journal'),
    ]
    add_styled_table(doc, ['Table', 'Records', 'Description'], schema_rows)

    doc.add_heading('Related Tables', level=2)
    related = [
        ('District', f'{len(all_districts)}', 'District master - links to AccountingFormats via AccountingFormatId'),
        ('Requisitions', '-', 'Purchase requisitions with account code assignments'),
        ('PO', '-', 'Purchase orders generated from approved requisitions'),
        ('Detail', '-', 'Requisition line items with pricing'),
        ('School', '-', 'Schools within districts for shipping/allocation'),
    ]
    add_styled_table(doc, ['Table', 'Active Records', 'Description'], related)

    doc.add_page_break()

    # ---- All Districts Reference ----
    doc.add_heading('Complete District Directory', level=1)
    doc.add_paragraph(f'All {len(all_districts)} active districts with their assigned accounting system.')

    all_dist_rows = []
    for d in sorted(all_districts, key=lambda x: x.get('Name', '')):
        all_dist_rows.append((
            d.get('Name', ''),
            d.get('City', ''),
            d.get('State', ''),
            (d.get('AccountingSystem') or 'None')[:35],
            str(d.get('AccountCount', 0)),
        ))

    add_styled_table(doc, ['District', 'City', 'ST', 'Accounting System', 'Accts'], all_dist_rows)

    filepath = os.path.join(OUTPUT_DIR, 'EDS_Accounting_Master_Summary.docx')
    doc.save(filepath)
    print(f"  Saved: {filepath}")
    return filepath


# ============================================================
# MAIN EXECUTION
# ============================================================

print("\n=== Generating Individual System Documents ===")
generated = []

for family in sorted(districts_by_family.keys()):
    districts = districts_by_family[family]
    print(f"\nGenerating: {family} ({len(districts)} districts)...")
    try:
        path = generate_system_doc(family, districts)
        generated.append(path)
    except Exception as e:
        print(f"  ERROR: {e}")

print("\n=== Generating Master Summary ===")
try:
    master_path = generate_master_summary()
    generated.append(master_path)
except Exception as e:
    print(f"  ERROR: {e}")

conn.close()

print(f"\n{'='*60}")
print(f"Documentation generation complete!")
print(f"Total documents generated: {len(generated)}")
print(f"Output directory: {os.path.abspath(OUTPUT_DIR)}")
for g in generated:
    size = os.path.getsize(g)
    print(f"  {os.path.basename(g)} ({size:,} bytes)")
